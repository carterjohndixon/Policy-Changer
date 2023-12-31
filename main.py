import os
from docx import Document

def replacer(document, fields: dict[str, str]):
    for paragraph in document.paragraphs:
        for field_key in fields.keys():
            if field_key in paragraph.text:
                paragraph.text = paragraph.text.replace(field_key, fields.get(field_key))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for field_key in fields.keys():
                    if field_key in cell.text:
                        cell.text = cell.text.replace(field_key, fields.get(field_key))

def does_field_exist_file(doc_path, field_to_replace) -> bool:
    document = Document(doc_path.strip())
    for paragraph in document.paragraphs:
        if field_to_replace in paragraph.text:
            return True

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if field_to_replace in cell.text:
                    return True

    return False

def does_field_exist_folder(doc_path, field_to_replace):
    for root, d, files in os.walk(doc_path):
        for file in files:
            file_path = os.path.join(root, file)
            for x in d:
                if x not in file_path:
                    if file_path.endswith('.docx'):
                        document = Document(file_path)
                        for paragraph in document.paragraphs:
                            if field_to_replace in paragraph.text:
                                return True

                        for table in document.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    if field_to_replace in cell.text:
                                        return True

    return False

def return_files_not_found_folders(doc_path, field_to_replace):
    files_not_found = []
    for root, d, files in os.walk(doc_path):
        for file in files:
            file_path = os.path.join(root, file)
            for x in d:
                if x not in file_path:
                    if file_path.endswith('.docx'):
                        if not does_field_exist_folder(file_path, field_to_replace):
                            files_not_found.append(file_path)

    return files_not_found

def find_and_replace_field_single_file(doc_path, modified_path_ques):
    document = Document(doc_path.strip())

    fields_dict: dict[str, str] = dict()

    while True:
        field_to_replace = input("Enter the field to replace (enter Q to quit or if done applying changes): ")

        if field_to_replace == 'Q' or field_to_replace == 'q':
            break

        is_field: bool = does_field_exist_file(doc_path, field_to_replace)
        while is_field == False:
            print(f"Field: {field_to_replace} does not exist in: {doc_path}")
            field_to_replace = input("Enter the field to replace (enter Q to quit or if done applying changes): ")
            is_field = does_field_exist_file(doc_path, field_to_replace)

        replacement_text = input(f"Enter what to change {field_to_replace} to: ")
        fields_dict[field_to_replace] = replacement_text

        replacer(document, fields_dict)

    if len(fields_dict) > 0:
        print("\nField changes:")
        for field, replacement in fields_dict.items():
            print(f"{field} -> {replacement}")

        confirm = input("\nDo you want to apply these changes? (Y/N): ")
        if confirm.upper() == 'Y':
            if modified_path_ques == "Y":
                doc_dir, doc_filename = os.path.split(doc_path)
                modified_doc_path = os.path.join(doc_dir + "/modified-" + doc_filename)
                document.save(modified_doc_path)
                print("Field replacement complete. Modified document saved as:", modified_doc_path)
            elif modified_path_ques == "N":
                doc_path = doc_path.strip()
                document.save(doc_path)
                print("Field replacement complete. Modified document saved as:", doc_path)
        else:
            find_and_replace_field_single_file(doc_path, modified_path_ques)
    else:
        print("No changes applied. Exiting program.")

def find_and_replace_field_folder(folder_directory, fields: dict[str, str]):
    for root, d, files in os.walk(folder_directory):
        for file in files:
            file_path = os.path.join(root, file)
            for x in d:
                if x not in file_path:
                    if file_path.endswith('.docx'):
                        doc = Document(file_path)
                        replacer(doc, fields)
                        doc.save(file_path)
                        print("Field replacement complete. Modified document saved as:", file_path)

def get_filename(file_path):
    return os.path.basename(file_path)

def get_files(fields_dict, doc_path):
    files_found = []

    for root, d, files in os.walk(doc_path):
        for file in files:
            file_path = os.path.join(root, file)
            for x in d:
                if x not in file_path:
                    if file_path.endswith('.docx'):
                        document = Document(file_path)
                        for paragraph in document.paragraphs:
                            for field, replacement in fields_dict.items():
                                if field in paragraph.text:
                                    if file_path not in files_found:
                                        files_found.append(file_path)
                                        break

                        for table in document.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for field, replacement in fields_dict.items():
                                        if field in cell.text:
                                            if file_path not in files_found:
                                                files_found.append(file_path)
                                                break
    return files_found

def apply_replacements(fields_dict, doc_path):
    if len(fields_dict) > 0:
        print("\nField changes:")
        for field, replacement in fields_dict.items():
            print(f"Field: {field} -> {replacement}")
            files_found = get_files(fields_dict, doc_path)
            for file_path in files_found:
                file_name = get_filename(file_path)
                print(f"  - File: {file_name}")

        confirm = input("\nDo you want to apply these changes? (Y/N): ")
        if confirm.upper() == 'Y':
            for root, d, files in os.walk(doc_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    for x in d:
                        if x not in file_path:
                            if file_path.endswith('.docx'):
                                doc = Document(file_path)
                                replacer(doc, fields_dict)
                                doc.save(file_path)
                                print("Field replacement complete. Modified document saved as:", file_path)
        elif confirm.upper() == 'N':
            replacement_text = input(f"Enter what to change {replacement} to (Press Q to quit): ")
            if replacement_text.upper() == 'Q' or replacement_text == 'q':
                print("No changes applied. Exiting program.")
                return
            else:
                for field, replacement in fields_dict.items():
                    fields_dict[field] = replacement_text
                apply_replacements(fields_dict, doc_path)
        else:
            print("No changes applied. Exiting program.")
    else:
        print("No changes applied. Exiting program.")

def replacements(doc_path):
    doc_path = doc_path.strip()
    fields_dict: dict[str, str] = dict()

    while True:
        field_to_replace = input("Enter the field to replace (enter Q to quit or if done applying changes): ")

        if field_to_replace == 'Q' or field_to_replace == 'q':
            break

        is_field = does_field_exist_folder(doc_path, field_to_replace)
        while not is_field:
            files_not_found = return_files_not_found_folders(doc_path, field_to_replace)
            for file in files_not_found:
                print(f"Field: {field_to_replace} does not exist in: {file}")
            field_to_replace = input("Enter the field to replace (enter Q to quit or if done applying changes): ")
            if field_to_replace == 'Q' or field_to_replace == 'q':
                break
            is_field = does_field_exist_folder(doc_path, field_to_replace)

        if field_to_replace == 'Q' or field_to_replace == 'q':
            break

        replacement_text = input(f"Enter what to change {field_to_replace} to: ")
        fields_dict[field_to_replace] = replacement_text

    apply_replacements(fields_dict, doc_path)

def check_path(doc_path) -> bool:
    doc_path = doc_path.strip()

    if os.path.exists(doc_path) and os.path.isdir(doc_path):
        return True
    else:
        print("Invalid path!")
        return False

def main():
    while True:
        file_or_folder_quest = input("Will this be a folder (Y) or file (N) (Y/N): ")
        if file_or_folder_quest in ['Y', 'N', 'y', 'n']:
            break
        else:
            file_or_folder_quest = input("Will this be a folder (Y) or file (N) (Y/N): ")
    doc_path = input("Enter the path to the Word document: ")
    valid_path: bool = check_path(doc_path)
    while valid_path == False:
        doc_path = input("Enter the path to the Word document: ")
        valid_path: bool = check_path(doc_path)
    if file_or_folder_quest == "N" or file_or_folder_quest == 'n':
        modified_path_ques = input("Do you want a new modified file (Y/N): ")
        find_and_replace_field_single_file(doc_path, modified_path_ques)
    elif file_or_folder_quest == "Y" or file_or_folder_quest == 'y':
        replacements(doc_path)

if __name__ == "__main__":
    main()